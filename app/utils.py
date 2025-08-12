
# import re
# from io import BytesIO

# def clean_text(text):
#     text = re.sub(r'<[^>]?>', '', text)
#     text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!$$$$,]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
#     text = re.sub(r'[^a-zA-Z0-9 ]', '', text)
#     text = re.sub(r'\s{2,}', ' ', text)
#     text = text.strip()
#     text = ' '.join(text.split())
#     return text

# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     try:
#         from pypdf import PdfReader
#     except ImportError:
#         raise ImportError("pypdf is required. Install with: pip install pypdf")
#     reader = PdfReader(BytesIO(file_bytes))
#     texts = []
#     for page in reader.pages:
#         txt = page.extract_text() or ""
#         texts.append(txt)
#     return "\n".join(texts)

# def extract_text_from_docx(file_bytes: bytes) -> str:
#     try:
#         import docx
#     except ImportError:
#         raise ImportError("python-docx is required. Install with: pip install python-docx")
#     doc = docx.Document(BytesIO(file_bytes))
#     return "\n".join([p.text for p in doc.paragraphs])

# def extract_resume_text(uploaded_file) -> str:
#     file_bytes = uploaded_file.read()
#     name = uploaded_file.name.lower()
#     if name.endswith(".pdf"):
#         raw = extract_text_from_pdf(file_bytes)
#     elif name.endswith(".docx"):
#         raw = extract_text_from_docx(file_bytes)
#     elif name.endswith(".txt"):
#         raw = file_bytes.decode("utf-8", errors="ignore")
#     else:
#         raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")
#     return raw


import re
from io import BytesIO

def clean_text(text):
    text = re.sub(r'<[^>]*>', '', text)
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!$&*,]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    text = re.sub(r'[^a-zA-Z0-9 ]', ' ', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()

def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx
    doc = docx.Document(BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs])

def extract_resume_text(uploaded_file) -> str:
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        raw = extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        raw = extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):
        raw = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")
    return raw
